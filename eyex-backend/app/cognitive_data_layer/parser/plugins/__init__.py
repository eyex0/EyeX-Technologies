from __future__ import annotations

from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

import pandas as pd
from charset_normalizer import detect

from app.cognitive_data_layer.parser.base import BaseParser, ParseResult


class CSVParser(BaseParser):
    name = "csv"
    supported_extensions = ["csv", "txt"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        ext = Path(source).suffix.lower().lstrip(".") if isinstance(source, (str, Path)) else ""
        return ext in self.supported_extensions

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        options = options or {}
        warnings: list[str] = []

        if isinstance(source, (str, Path)):
            with open(source, "rb") as f:
                raw_bytes = f.read()
        else:
            raw_bytes = source

        encoding = options.get("encoding")
        if not encoding:
            detection = detect(raw_bytes)
            encoding = detection.get("encoding", "utf-8") or "utf-8"
            warnings.append(f"Detected encoding: {encoding}")

        try:
            text = raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            text = raw_bytes.decode("utf-8", errors="replace")
            warnings.append("Fallback decode with replacement characters")

        df = pd.read_csv(
            StringIO(text),
            sep=options.get("sep", ","),
            engine="python",
            on_bad_lines="warn",
            dtype=str,
            keep_default_na=True,
        )
        df = _clean_dataframe(df)

        return ParseResult(
            raw_data={"sheets": [{"name": "default", "data": df}]},
            format="csv",
            metadata={"encoding": encoding, "rows": len(df), "columns": len(df.columns)},
            warnings=warnings,
        )


class ExcelParser(BaseParser):
    name = "excel"
    supported_extensions = ["xlsx", "xls"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        if isinstance(source, (str, Path)):
            ext = Path(source).suffix.lower().lstrip(".")
            return ext in self.supported_extensions
        from io import BytesIO

        if isinstance(source, BytesIO):
            source.seek(0)
            header = source.read(4)
            source.seek(0)
            return header == b"PK\x03\x04"
        if isinstance(source, bytes):
            return source.startswith(b"PK")
        return False

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        options = options or {}
        warnings: list[str] = []

        if isinstance(source, (str, Path)):
            file = source
        else:
            file = BytesIO(source)

        xl = pd.ExcelFile(file)
        sheets: list[dict[str, Any]] = []
        for sheet_name in xl.sheet_names:
            try:
                df = pd.read_excel(file, sheet_name=sheet_name, dtype=str, keep_default_na=True)
                df = _clean_dataframe(df)
                sheets.append({"name": sheet_name, "data": df})
            except Exception as exc:  # noqa: BLE001
                warnings.append(f"Could not read sheet '{sheet_name}': {exc}")

        return ParseResult(
            raw_data={"sheets": sheets},
            format="excel",
            metadata={
                "sheet_names": xl.sheet_names,
                "rows": sum(s["data"].shape[0] for s in sheets),
                "columns": sum(s["data"].shape[1] for s in sheets),
            },
            warnings=warnings,
        )


class JSONParser(BaseParser):
    name = "json"
    supported_extensions = ["json"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        ext = Path(source).suffix.lower().lstrip(".") if isinstance(source, (str, Path)) else ""
        return ext == "json"

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        import json

        if isinstance(source, (str, Path)):
            with open(source, encoding="utf-8") as f:
                data = json.load(f)
        else:
            data = json.loads(source.decode("utf-8"))

        sheets = [{"name": "default", "data": _json_to_dataframe(data)}]
        return ParseResult(
            raw_data={"sheets": sheets},
            format="json",
            metadata={"rows": len(sheets[0]["data"]), "columns": len(sheets[0]["data"].columns)},
            warnings=[],
        )


class XMLParser(BaseParser):
    name = "xml"
    supported_extensions = ["xml"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        ext = Path(source).suffix.lower().lstrip(".") if isinstance(source, (str, Path)) else ""
        return ext == "xml"

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        import xml.etree.ElementTree as ET

        if isinstance(source, (str, Path)):
            with open(source, "rb") as f:
                raw_bytes = f.read()
        else:
            raw_bytes = source

        root = ET.fromstring(raw_bytes)
        records = _xml_to_records(root)
        df = pd.DataFrame(records)
        df = _clean_dataframe(df)

        return ParseResult(
            raw_data={"sheets": [{"name": "default", "data": df}]},
            format="xml",
            metadata={"rows": len(df), "columns": len(df.columns)},
            warnings=[],
        )


class SQLParser(BaseParser):
    name = "sql"
    supported_extensions = ["sql"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        ext = Path(source).suffix.lower().lstrip(".") if isinstance(source, (str, Path)) else ""
        return ext == "sql"

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        from sqlalchemy import create_engine, text

        if isinstance(source, (str, Path)):
            connection_string = str(source)
        else:
            connection_string = source.decode("utf-8")

        engine = create_engine(connection_string)
        query = options.get("query", "SELECT * FROM information_schema.tables")
        with engine.connect() as conn:
            df = pd.read_sql(text(query), conn)
        df = _clean_dataframe(df)

        return ParseResult(
            raw_data={"sheets": [{"name": "sql_result", "data": df}]},
            format="sql",
            metadata={"rows": len(df), "columns": len(df.columns)},
            warnings=[],
        )


class APIParser(BaseParser):
    name = "api"
    supported_extensions: list[str] = []

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        return isinstance(source, str) and (
            source.startswith("http://") or source.startswith("https://")
        )

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        import httpx

        options = options or {}
        if not isinstance(source, str):
            raise ValueError("API parser requires a URL string")

        async with httpx.AsyncClient() as client:
            response = await client.get(source, headers=options.get("headers", {}))
            response.raise_for_status()
            data = response.json()

        df = _json_to_dataframe(data)
        return ParseResult(
            raw_data={"sheets": [{"name": "api_response", "data": df}]},
            format="api",
            metadata={"rows": len(df), "columns": len(df.columns)},
            warnings=[],
        )


class PDFParser(BaseParser):
    name = "pdf"
    supported_extensions = ["pdf"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        ext = Path(source).suffix.lower().lstrip(".") if isinstance(source, (str, Path)) else ""
        return ext == "pdf"

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        import pdfplumber

        if isinstance(source, (str, Path)):
            file = source
        else:
            file = BytesIO(source)

        warnings: list[str] = []
        tables: list[dict[str, Any]] = []
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        df = _clean_dataframe(df)
                        tables.append({"name": f"page_{i + 1}_table", "data": df})

        if not tables:
            warnings.append("No tables detected in PDF")

        return ParseResult(
            raw_data={"sheets": tables},
            format="pdf",
            metadata={"pages": len(pdf.pages), "tables": len(tables)},
            warnings=warnings,
        )


class DOCXParser(BaseParser):
    name = "docx"
    supported_extensions = ["docx"]

    def can_parse(self, source: str | Path | bytes, hint: str | None = None) -> bool:
        ext = Path(source).suffix.lower().lstrip(".") if isinstance(source, (str, Path)) else ""
        return ext == "docx"

    async def parse(self, source: str | Path | bytes, options: dict | None = None) -> ParseResult:
        from docx import Document

        if isinstance(source, (str, Path)):
            doc = Document(source)
        else:
            doc = Document(BytesIO(source))

        tables: list[dict[str, Any]] = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            if rows:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                df = _clean_dataframe(df)
                tables.append({"name": f"table_{i + 1}", "data": df})

        return ParseResult(
            raw_data={"sheets": tables},
            format="docx",
            metadata={"tables": len(tables)},
            warnings=[],
        )


def _clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all").reset_index(drop=True)
    df = df.loc[:, ~df.columns.str.contains("^Unnamed")]
    return df


def _json_to_dataframe(data: Any) -> pd.DataFrame:
    if isinstance(data, list):
        return pd.json_normalize(data)
    if isinstance(data, dict):
        return pd.json_normalize(data)
    return pd.DataFrame([data])


def _xml_to_records(element: Any) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for child in element:
        if len(child) > 0:
            record = {c.tag: c.text or "" for c in child}
            records.append(record)
    return records


def register_default_parsers() -> None:
    from app.cognitive_data_layer.parser.registry import register_parser

    register_parser(CSVParser())
    register_parser(ExcelParser())
    register_parser(JSONParser())
    register_parser(XMLParser())
    register_parser(SQLParser())
    register_parser(APIParser())
    register_parser(PDFParser())
    register_parser(DOCXParser())
